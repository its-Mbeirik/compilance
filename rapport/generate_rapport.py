#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Générateur du rapport PFE ConformIA — ISCAE 2025-2026"""

import os, io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RAPPORT_DIR = r"C:\Users\bouda\Desktop\iscae\s6\pfe2026\assistant_conformite\rapport"
OUTPUT = os.path.join(RAPPORT_DIR, "Rapport_PFE_ConformIA.docx")

def img(name):
    return os.path.join(RAPPORT_DIR, name)

def svg_to_bytes(path):
    try:
        import cairosvg
        return cairosvg.svg2png(url=path, output_width=1400)
    except Exception:
        pass
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        d = svg2rlg(path)
        if d:
            buf = io.BytesIO()
            renderPM.drawToFile(d, buf, fmt="PNG", dpi=150)
            buf.seek(0)
            return buf.read()
    except Exception:
        pass
    return None

def insert_image(doc, filename, width=Inches(5.8)):
    path = img(filename)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    ok = False
    if filename.endswith('.svg'):
        data = svg_to_bytes(path)
        if data:
            run.add_picture(io.BytesIO(data), width=width)
            ok = True
    elif os.path.exists(path):
        try:
            run.add_picture(path, width=width)
            ok = True
        except Exception:
            pass
    if not ok:
        para.clear()
        r = para.add_run(f"[Insérer ici : {filename}]")
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    return para

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(14)
    return p

def figure(doc, filename, cap_text, width=Inches(5.8)):
    insert_image(doc, filename, width)
    caption(doc, cap_text)

def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    hd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return hd

def txt(doc, text, bold=False, italic=False, size=12,
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, before=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def bul(doc, text, size=12):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

def centered(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    return p

def add_table_row(table, cells, bold_first=False, shade=None):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = val
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(11)
                if bold_first and i == 0:
                    run.bold = True
    return row

# ─────────────────────────────────────────────
def build():
    doc = Document()

    # Margins
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    # Default font
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # ──────────────────────────────────────────
    # PAGE DE GARDE
    # ──────────────────────────────────────────
    for _ in range(2):
        doc.add_paragraph()

    figure(doc, 'logo-iscae.png', '', width=Inches(2.2))

    centered(doc,
        "INSTITUT SUPERIEUR DE COMPTABILITE ET D'ADMINISTRATION DES ENTREPRISES",
        size=13, bold=True)
    centered(doc, "Nouakchott — Mauritanie", size=12)

    for _ in range(2):
        doc.add_paragraph()

    p_titre = doc.add_paragraph()
    p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_titre.add_run("PROJET FIN D'ETUDES")
    r.font.size = Pt(16)
    r.bold = True

    doc.add_paragraph()
    centered(doc, "Filière :", size=12, bold=True)
    centered(doc, "Développement Informatique (DI)", size=12)

    doc.add_paragraph()
    centered(doc, "Thème :", size=12, bold=True)

    p_theme = doc.add_paragraph()
    p_theme.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_theme.add_run(
        "Conception et développement d'un assistant intelligent\n"
        "de conformité contractuelle basé sur l'IA générative\n"
        "— ConformIA —"
    )
    r.font.size = Pt(14)
    r.bold = True

    for _ in range(2):
        doc.add_paragraph()

    centered(doc, "Réalisé par :", size=12, bold=True)
    centered(doc, "Ahmed Bouda                          I20538", size=12)
    centered(doc, "Oumneaicha Eljilli                      I20438", size=12)
    centered(doc, "Naha Hamadi                            I20650", size=12)

    doc.add_paragraph()
    centered(doc, "Encadré par :", size=12, bold=True)
    centered(doc, "DR Mohamed Jibril", size=12)

    for _ in range(2):
        doc.add_paragraph()

    centered(doc, "Année universitaire : 2025-2026", size=12, bold=True)

    doc.add_page_break()

    # ──────────────────────────────────────────
    # DEDICACE
    # ──────────────────────────────────────────
    h(doc, "Dédicace", level=1)
    doc.add_paragraph()

    txt(doc,
        "À nos familles, pour leur amour inconditionnel, leur patience infinie et leurs sacrifices "
        "silencieux qui ont rendu possible chacune de nos réussites. Ce travail est le leur autant que le nôtre.")
    doc.add_paragraph()
    txt(doc,
        "À nos enseignants et encadrants, en particulier DR Mohamed Jibril, pour la richesse de son "
        "enseignement, sa disponibilité et ses précieux conseils tout au long de ce projet.")
    doc.add_paragraph()
    txt(doc,
        "À nos amis et collègues, pour leur soutien moral, leur encouragement constant et leur présence "
        "réconfortante dans les moments de doute.")
    doc.add_paragraph()
    txt(doc,
        "À toutes celles et ceux qui, de près ou de loin, ont contribué à la réussite de ce projet "
        "de fin d'études.",
        italic=True)

    doc.add_page_break()

    # ──────────────────────────────────────────
    # REMERCIEMENTS
    # ──────────────────────────────────────────
    h(doc, "Remerciements", level=1)
    doc.add_paragraph()

    txt(doc,
        "Avant tout, nous rendons grâce à Dieu Tout-Puissant pour nous avoir accordé la force, "
        "la santé et la persévérance nécessaires à l'accomplissement de ce projet.")
    doc.add_paragraph()
    txt(doc,
        "Nous tenons à exprimer notre profonde gratitude à notre encadrant, DR Mohamed Jibril, "
        "pour sa disponibilité remarquable, ses conseils éclairés et son suivi rigoureux tout au long "
        "de ce projet de fin d'études. Sa bienveillance et son expertise ont été déterminantes dans "
        "la réussite de ce travail.")
    doc.add_paragraph()
    txt(doc,
        "Nous remercions chaleureusement l'ensemble du corps enseignant de l'Institut Supérieur de "
        "Comptabilité et d'Administration des Entreprises (ISCAE) pour la qualité de la formation "
        "dispensée et l'accompagnement pédagogique dont nous avons bénéficié tout au long de notre cursus.")
    doc.add_paragraph()
    txt(doc,
        "Nos remerciements vont également à la direction de l'ISCAE pour les moyens mis à notre "
        "disposition et pour l'environnement propice à l'apprentissage et à l'innovation qu'elle a su créer.")
    doc.add_paragraph()
    txt(doc,
        "Enfin, nous remercions sincèrement nos familles et nos proches pour leur soutien indéfectible "
        "et leurs encouragements tout au long de notre parcours universitaire.")

    doc.add_page_break()

    # ──────────────────────────────────────────
    # RESUME
    # ──────────────────────────────────────────
    h(doc, "Résumé", level=1)
    doc.add_paragraph()

    txt(doc,
        "Dans un contexte où la conformité contractuelle représente un enjeu stratégique majeur pour les "
        "entreprises mauritaniennes, ce projet de fin d'études s'inscrit dans la conception et le développement "
        "d'une application web intelligente dédiée à l'analyse automatique de la conformité des contrats "
        "de travail avec le droit mauritanien.")
    doc.add_paragraph()
    txt(doc,
        "La problématique centrale réside dans la difficulté pour les entreprises, les juristes et les "
        "professionnels des ressources humaines de vérifier manuellement la conformité de leurs contrats "
        "avec le Code du Travail mauritanien (Loi N° 2004-017), le Code des Obligations et des Contrats "
        "(Ordonnance N° 89-126), le Code du Commerce (Loi N° 2000-05) et les conventions collectives. "
        "Ce processus, chronophage et source d'erreurs, expose les entreprises à des risques juridiques "
        "et financiers significatifs.")
    doc.add_paragraph()
    txt(doc,
        "La solution développée, ConformIA, repose sur une architecture moderne alliant Next.js 15 et "
        "React 19 pour le frontend, FastAPI pour le backend, et LangGraph pour l'orchestration d'un "
        "pipeline d'intelligence artificielle composé de trois agents spécialisés : un extracteur de "
        "clauses, un récupérateur sémantique basé sur les embeddings (BAAI/bge-m3 + pgvector) et un "
        "vérificateur doté d'un mécanisme de garde contre les hallucinations (Citation Guard). "
        "Le corpus juridique indexé comprend 2 176 articles issus des principaux textes législatifs mauritaniens.")
    doc.add_paragraph()
    txt(doc,
        "Les résultats obtenus permettent de livrer une application complète offrant : l'analyse de "
        "conformité en temps réel avec verdicts détaillés (CONFORME, NON_CONFORME, EXIGE_REVUE), "
        "la génération de rapports PDF, la correction automatique des contrats, un assistant juridique "
        "conversationnel bilingue (français/arabe) avec saisie vocale, ainsi qu'un système de gestion "
        "des utilisateurs à trois niveaux avec vérification d'email et réinitialisation de mot de passe.")

    doc.add_paragraph()
    p_mots = doc.add_paragraph()
    p_mots.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p_mots.add_run("Mots-clés : ")
    r.bold = True
    r.font.size = Pt(12)
    r2 = p_mots.add_run(
        "Intelligence artificielle, conformité contractuelle, droit du travail mauritanien, "
        "RAG, LangGraph, LLM, pgvector, Next.js, FastAPI."
    )
    r2.italic = True
    r2.font.size = Pt(12)

    doc.add_page_break()

    # ──────────────────────────────────────────
    # TABLE DES MATIERES
    # ──────────────────────────────────────────
    h(doc, "Table des matières", level=1)
    doc.add_paragraph()

    toc_items = [
        ("Dédicace", "2"),
        ("Remerciements", "3"),
        ("Résumé", "4"),
        ("Table des matières", "5"),
        ("Liste des acronymes", "6"),
        ("Liste des figures", "7"),
        ("Introduction générale", "8"),
        ("Chapitre I : Présentation de l'institution et du contexte", "10"),
        ("    1.1  Présentation de l'ISCAE", "10"),
        ("    1.2  Domaines d'intervention et formation en informatique", "11"),
        ("    1.3  Cadre juridique mauritanien", "11"),
        ("    1.4  Conclusion", "12"),
        ("Chapitre II : Contexte et problématique du projet", "13"),
        ("    2.1  Contexte général", "13"),
        ("    2.2  Problématique", "13"),
        ("    2.3  Objectifs et périmètre du projet", "14"),
        ("    2.4  Méthodologie de travail adoptée", "15"),
        ("    2.5  Conclusion", "16"),
        ("Chapitre III : Méthodologie et aspects techniques", "17"),
        ("    3.1  Méthodologie Scrum", "17"),
        ("    3.2  Architecture technique détaillée", "18"),
        ("    3.3  Spécifications fonctionnelles", "19"),
        ("    3.4  Diagrammes fonctionnels", "20"),
        ("    3.5  Conclusion", "24"),
        ("Chapitre IV : Développement et réalisation", "25"),
        ("    4.1  Choix des technologies", "25"),
        ("    4.2  Mise en place du backend (FastAPI + LangGraph)", "27"),
        ("    4.3  Développement du frontend (Next.js)", "29"),
        ("    4.4  Sécurité et gestion des utilisateurs", "30"),
        ("    4.5  Corpus juridique et système RAG", "31"),
        ("    4.6  Présentation de l'interface utilisateur", "32"),
        ("    4.7  Conclusion", "38"),
        ("Conclusion générale", "39"),
        ("Bibliographie", "42"),
    ]

    for item, page in toc_items:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.space_after = Pt(2)
        tab_stops = p_toc.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(14), 2)  # right align tab
        r = p_toc.add_run(item + "\t" + page)
        r.font.size = Pt(11)
        if not item.startswith("    "):
            r.bold = True

    doc.add_page_break()

    # ──────────────────────────────────────────
    # LISTE DES ACRONYMES
    # ──────────────────────────────────────────
    h(doc, "Liste des acronymes", level=1)
    doc.add_paragraph()

    acronyms = [
        ("API",     "Application Programming Interface",           "Interface de programmation d'application"),
        ("BGE",     "BAAI General Embedding",                      "Modèle d'embedding de BAAI"),
        ("CDI",     "Contrat à Durée Indéterminée",               "Open-ended employment contract"),
        ("CDD",     "Contrat à Durée Déterminée",                 "Fixed-term employment contract"),
        ("CI/CD",   "Continuous Integration / Continuous Deployment", "Intégration et déploiement continus"),
        ("COC",     "Code des Obligations et des Contrats",        "Code of Obligations and Contracts"),
        ("CORS",    "Cross-Origin Resource Sharing",               "Partage de ressources entre origines"),
        ("CSS",     "Cascading Style Sheets",                      "Feuilles de style en cascade"),
        ("DOCX",    "Document XML Format",                         "Format de document Microsoft Word"),
        ("IA",      "Intelligence Artificielle",                   "Artificial Intelligence"),
        ("JWT",     "JSON Web Token",                              "Jeton d'authentification JSON"),
        ("LLM",     "Large Language Model",                        "Grand modèle de langage"),
        ("OIT",     "Organisation Internationale du Travail",      "International Labour Organization"),
        ("ORM",     "Object-Relational Mapping",                   "Mapping objet-relationnel"),
        ("PDF",     "Portable Document Format",                    "Format de document portable"),
        ("RAG",     "Retrieval-Augmented Generation",              "Génération augmentée par récupération"),
        ("REST",    "Representational State Transfer",             "Transfert d'état représentatif"),
        ("SQL",     "Structured Query Language",                   "Langage de requête structuré"),
        ("UI",      "User Interface",                              "Interface utilisateur"),
        ("UX",      "User Experience",                             "Expérience utilisateur"),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = "Acronyme"
    hdr[1].text = "Signification (EN)"
    hdr[2].text = "Signification (FR)"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)

    for acr, en, fr in acronyms:
        row = table.add_row()
        row.cells[0].text = acr
        row.cells[1].text = en
        row.cells[2].text = fr
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ──────────────────────────────────────────
    # LISTE DES FIGURES
    # ──────────────────────────────────────────
    h(doc, "Liste des figures", level=1)
    doc.add_paragraph()

    figures_list = [
        ("Figure 1",  "Architecture générale de ConformIA"),
        ("Figure 2",  "Diagramme des cas d'utilisation"),
        ("Figure 3",  "Diagramme de classes"),
        ("Figure 4",  "Diagramme de séquence — Analyse de conformité"),
        ("Figure 5",  "Page de connexion"),
        ("Figure 6",  "Page d'inscription"),
        ("Figure 7",  "Page de vérification d'adresse email"),
        ("Figure 8",  "Page « Mot de passe oublié »"),
        ("Figure 9",  "Page de réinitialisation du mot de passe"),
        ("Figure 10", "Interface principale de l'assistant juridique"),
        ("Figure 11", "Résultats d'analyse de conformité contractuelle"),
        ("Figure 12", "Génération et téléchargement de document corrigé"),
        ("Figure 13", "Tableau de bord administrateur"),
        ("Figure 14", "Gestion des approbations en attente"),
        ("Figure 15", "Gestion des sous-utilisateurs"),
        ("Figure 16", "Paramètres du compte utilisateur"),
    ]

    fig_table = doc.add_table(rows=1, cols=2)
    fig_table.style = 'Table Grid'
    fhdr = fig_table.rows[0].cells
    fhdr[0].text = "N°"
    fhdr[1].text = "Titre de la figure"
    for cell in fhdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(11)

    for num, title in figures_list:
        row = fig_table.add_row()
        row.cells[0].text = num
        row.cells[1].text = title
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ──────────────────────────────────────────
    # INTRODUCTION GENERALE
    # ──────────────────────────────────────────
    h(doc, "Introduction générale", level=1)
    doc.add_paragraph()

    txt(doc,
        "L'essor de l'intelligence artificielle générative ouvre des perspectives inédites dans le domaine "
        "du droit et de la conformité réglementaire. En Mauritanie, comme dans l'ensemble des pays en "
        "développement, la mise en conformité des contrats de travail avec la législation nationale "
        "constitue un défi quotidien pour les entreprises, les juristes et les services des ressources humaines. "
        "Les textes législatifs, nombreux et complexes, sont souvent méconnus ou mal appliqués, exposant "
        "les employeurs à des litiges coûteux et les employés à des violations de leurs droits fondamentaux.")
    doc.add_paragraph()
    txt(doc,
        "Le présent travail s'inscrit dans le cadre du projet de fin d'études de la filière Développement "
        "Informatique à l'Institut Supérieur de Comptabilité et d'Administration des Entreprises (ISCAE) "
        "de Nouakchott. Il a pour objet la conception et le développement d'une application web intelligente, "
        "ConformIA, dédiée à l'analyse automatique de la conformité des contrats de travail avec le droit "
        "mauritanien, en s'appuyant sur les technologies les plus récentes d'intelligence artificielle et "
        "de traitement du langage naturel.")
    doc.add_paragraph()
    txt(doc,
        "La problématique centrale à laquelle répond ce projet peut se formuler ainsi : comment concevoir "
        "un système intelligent capable de vérifier automatiquement et de manière fiable la conformité "
        "d'un contrat de travail avec l'ensemble du corpus juridique mauritanien, en fournissant des "
        "verdicts argumentés, des recommandations concrètes et en évitant les erreurs d'hallucination "
        "propres aux modèles de langage ?")
    doc.add_paragraph()
    txt(doc,
        "Pour répondre à cette problématique, nous avons conçu un pipeline d'analyse en trois étapes — "
        "extraction, récupération sémantique et vérification — orchestré par LangGraph, et appuyé sur "
        "un corpus de 2 176 articles juridiques mauritaniens indexés en base de données vectorielle. "
        "L'interface utilisateur, développée avec Next.js 15, offre une expérience bilingue "
        "(français et arabe) avec saisie vocale, et intègre des fonctionnalités avancées de génération "
        "et de correction de documents contractuels.")
    doc.add_paragraph()
    txt(doc, "Ce rapport est structuré comme suit :")
    bul(doc, "Le Chapitre I présente l'institution d'accueil (ISCAE) et le cadre juridique mauritanien dans lequel s'inscrit le projet.")
    bul(doc, "Le Chapitre II expose la problématique, les objectifs du projet et la méthodologie de travail adoptée.")
    bul(doc, "Le Chapitre III décrit l'architecture technique, les spécifications fonctionnelles et les diagrammes de conception.")
    bul(doc, "Le Chapitre IV présente le processus de développement, les choix technologiques et les interfaces réalisées.")
    bul(doc, "La conclusion générale revient sur les résultats obtenus, les difficultés rencontrées et les perspectives d'amélioration.")

    doc.add_page_break()

    # ──────────────────────────────────────────
    # CHAPITRE I
    # ──────────────────────────────────────────
    h(doc, "Chapitre I : Présentation de l'institution et du contexte", level=1)

    h(doc, "1.1  Identification de l'institution", level=2)
    txt(doc,
        "L'Institut Supérieur de Comptabilité et d'Administration des Entreprises (ISCAE) est un "
        "établissement d'enseignement supérieur public mauritanien, situé à Nouakchott. Fondé dans "
        "le cadre de la politique nationale de promotion de l'enseignement supérieur et de la formation "
        "des cadres qualifiés, l'ISCAE constitue un pilier de la formation en gestion, comptabilité, "
        "finance et développement informatique en Mauritanie.")
    doc.add_paragraph()
    txt(doc,
        "L'établissement dispense des formations de niveau Licence et Master dans plusieurs filières, "
        "dont la filière Développement Informatique (DI) au sein de laquelle s'inscrit ce projet. "
        "Sa mission principale est de former des techniciens et cadres supérieurs capables de répondre "
        "aux besoins croissants du marché mauritanien en compétences numériques et managériales.")

    h(doc, "1.2  Formation en Développement Informatique", level=2)
    txt(doc,
        "La filière Développement Informatique de l'ISCAE forme des ingénieurs logiciels maîtrisant "
        "les technologies web modernes, les systèmes d'information, les bases de données et l'intelligence "
        "artificielle. Les étudiants y acquièrent des compétences techniques solides en programmation "
        "(Python, JavaScript, SQL), en architecture logicielle et en gestion de projets informatiques.")
    doc.add_paragraph()
    txt(doc,
        "Le Projet de Fin d'Études (PFE) constitue l'aboutissement du cursus : il permet aux étudiants "
        "de mettre en pratique l'ensemble des connaissances acquises au cours de leur formation en "
        "réalisant un projet concret, innovant et techniquement avancé. C'est dans ce cadre que "
        "s'inscrit ConformIA.")

    h(doc, "1.3  Cadre juridique mauritanien", level=2)
    txt(doc,
        "La Mauritanie dispose d'un corpus législatif en matière de droit du travail et des contrats "
        "relativement complet, bien que peu connu du grand public et des petites entreprises. Les "
        "principaux textes de référence dans le domaine de la conformité contractuelle sont :")
    bul(doc, "Le Code du Travail mauritanien (Loi N° 2004-017) : texte fondamental régissant les relations de travail, comprenant 450 articles couvrant les contrats, les congés, les salaires, les conditions de travail et la protection des travailleurs.")
    bul(doc, "Le Code des Obligations et des Contrats (Ordonnance N° 89-126, modifiée par la Loi N° 2001-31) : 1 181 articles définissant les règles générales applicables à tous les contrats en droit mauritanien.")
    bul(doc, "Le Code du Commerce (Loi N° 2000-05) : 424 articles encadrant les pratiques commerciales et les contrats d'affaires.")
    bul(doc, "La Convention Collective Générale du Travail (UNICEMA/UTM) : 71 articles précisant les droits et obligations des parties dans les relations collectives de travail.")
    bul(doc, "Les Conventions Internationales du Travail (OIT) ratifiées par la Mauritanie : 50 articles établissant des standards internationaux en matière de droit du travail.")
    doc.add_paragraph()
    txt(doc,
        "Ce corpus, riche de 2 176 articles au total, constitue la base légale que ConformIA indexe "
        "et exploite pour analyser la conformité des contrats soumis par les utilisateurs.")

    h(doc, "1.4  Conclusion", level=2)
    txt(doc,
        "Ce premier chapitre a présenté l'ISCAE comme institution de formation et le cadre juridique "
        "mauritanien qui constitue le cœur du projet ConformIA. La richesse et la complexité du corpus "
        "législatif mauritanien, peu accessible au grand public, justifient pleinement le développement "
        "d'un outil d'assistance intelligent permettant d'en faciliter l'exploitation dans le cadre "
        "de la vérification contractuelle.")

    doc.add_page_break()

    # ──────────────────────────────────────────
    # CHAPITRE II
    # ──────────────────────────────────────────
    h(doc, "Chapitre II : Contexte et problématique du projet", level=1)

    h(doc, "2.1  Contexte général", level=2)
    txt(doc,
        "Dans un contexte de transformation numérique accélérée, les entreprises mauritaniennes font "
        "face à une double pression : d'une part, la nécessité de se conformer à une législation du "
        "travail de plus en plus stricte ; d'autre part, le manque de ressources juridiques internes "
        "capables d'assurer cette conformité au quotidien. Les petites et moyennes entreprises, qui "
        "représentent la majorité du tissu économique mauritanien, n'ont généralement pas les moyens "
        "de s'offrir les services d'un juriste spécialisé pour chaque contrat.")
    doc.add_paragraph()
    txt(doc,
        "Par ailleurs, les récentes avancées en intelligence artificielle, notamment les modèles de "
        "langage large (LLM) et les techniques de Retrieval-Augmented Generation (RAG), ouvrent de "
        "nouvelles perspectives pour automatiser des tâches juridiques complexes tout en maintenant "
        "un niveau de fiabilité élevé. C'est dans cette convergence entre besoin métier et opportunité "
        "technologique que s'inscrit le projet ConformIA.")

    h(doc, "2.2  Problématique", level=2)
    txt(doc,
        "Avant le développement de ConformIA, la vérification de la conformité d'un contrat de travail "
        "avec le droit mauritanien reposait entièrement sur une expertise humaine, avec les limitations "
        "suivantes :")
    bul(doc, "Processus chronophage : la vérification manuelle d'un contrat par rapport à l'ensemble des textes législatifs applicables peut prendre plusieurs heures.")
    bul(doc, "Risque d'omissions : la quantité et la complexité des textes rendent difficile une vérification exhaustive sans outil dédié.")
    bul(doc, "Inaccessibilité : les petites structures n'ont pas accès à des consultants juridiques spécialisés en droit du travail mauritanien.")
    bul(doc, "Absence de traçabilité : sans outil formel, il est difficile de documenter et d'archiver les analyses effectuées.")
    doc.add_paragraph()
    txt(doc,
        "La problématique principale à laquelle répond ce projet peut donc se formuler ainsi : "
        "Comment concevoir une solution applicative intelligente, fiable et accessible permettant "
        "d'analyser automatiquement la conformité d'un contrat de travail avec le corpus juridique "
        "mauritanien, tout en évitant les erreurs d'hallucination propres aux modèles de langage ?",
        bold=False, italic=True)

    h(doc, "2.3  Objectifs et périmètre du projet", level=2)
    txt(doc, "L'objectif principal du projet est de :", bold=True)
    doc.add_paragraph()
    txt(doc,
        "Développer une application web intelligente, ConformIA, permettant l'analyse automatique "
        "de la conformité des contrats de travail avec le droit mauritanien, la génération de rapports "
        "détaillés, la correction automatique des contrats non conformes, et l'assistance juridique "
        "conversationnelle bilingue.")
    doc.add_paragraph()
    txt(doc, "Objectifs spécifiques :", bold=True)
    bul(doc, "Constituer et indexer un corpus de 2 176 articles juridiques mauritaniens dans une base de données vectorielle (PostgreSQL + pgvector).")
    bul(doc, "Développer un pipeline d'analyse à trois agents (extracteur, récupérateur, vérificateur) orchestré par LangGraph.")
    bul(doc, "Implémenter un mécanisme de garde contre les hallucinations (Citation Guard) pour garantir la fiabilité des citations légales.")
    bul(doc, "Créer une interface utilisateur moderne, bilingue (français/arabe) avec saisie vocale, développée avec Next.js 15.")
    bul(doc, "Mettre en place un système de gestion des utilisateurs à trois niveaux (Admin, Utilisateur, Sous-utilisateur) avec vérification d'email et réinitialisation de mot de passe.")
    bul(doc, "Déployer l'application via Docker Compose pour une portabilité et une reproductibilité optimales.")
    doc.add_paragraph()
    txt(doc, "Périmètre du projet :", bold=True)
    txt(doc,
        "Le périmètre fonctionnel couvre l'analyse de contrats de travail de droit mauritanien "
        "(CDI, CDD, CTT, Stage), la génération de rapports de conformité, la correction automatique, "
        "le chat juridique et la gestion complète des utilisateurs. Les fonctionnalités annexes telles "
        "que l'intégration avec des systèmes tiers de gestion RH ou la prise en charge de corpus "
        "législatifs d'autres pays sont envisagées comme perspectives futures.")

    h(doc, "2.4  Méthodologie de travail adoptée", level=2)
    txt(doc,
        "Le projet a été conduit selon la méthodologie agile Scrum, particulièrement adaptée aux "
        "projets de développement logiciel innovants où les besoins évoluent au fil des itérations. "
        "Cette approche a permis de structurer le travail en sprints de deux semaines, chacun "
        "donnant lieu à une livraison partielle et à une révision avec l'encadrant.")
    doc.add_paragraph()
    txt(doc, "Les grandes phases du projet ont été :", bold=True)
    bul(doc, "Sprint 1 — Analyse et conception : étude du corpus juridique, modélisation de la base de données, définition de l'architecture technique, mise en place de l'environnement Docker.")
    bul(doc, "Sprint 2 — Ingestion du corpus : développement du pipeline d'ingestion des 2 176 articles, génération des embeddings BGE-M3, indexation dans pgvector.")
    bul(doc, "Sprint 3 — Pipeline LangGraph : développement des agents extracteur, récupérateur et vérificateur, implémentation du Citation Guard.")
    bul(doc, "Sprint 4 — API et authentification : développement des routes FastAPI, système JWT, gestion des utilisateurs, vérification d'email, réinitialisation de mot de passe.")
    bul(doc, "Sprint 5 — Frontend : développement de l'interface Next.js, intégration bilingue français/arabe, saisie vocale, tableau de bord administrateur.")
    bul(doc, "Sprint 6 — Tests, intégration et finalisation : tests unitaires, CI/CD, génération de documents DOCX, rapports PDF, polish de l'interface.")

    h(doc, "2.5  Conclusion", level=2)
    txt(doc,
        "Ce chapitre a exposé le contexte général de la conformité contractuelle en Mauritanie, "
        "la problématique à laquelle répond ConformIA, ses objectifs précis et la méthodologie "
        "Scrum adoptée pour mener à bien le projet. La structure itérative de la méthode agile "
        "a permis d'adapter continuellement la solution aux besoins identifiés, tout en maintenant "
        "un niveau de qualité élevé à chaque étape du développement.")

    doc.add_page_break()

    # ──────────────────────────────────────────
    # CHAPITRE III
    # ──────────────────────────────────────────
    h(doc, "Chapitre III : Méthodologie et aspects techniques", level=1)

    h(doc, "3.1  Méthodologie Scrum", level=2)
    txt(doc,
        "La méthode Scrum adoptée s'est articulée autour d'un cycle itératif structuré. Chaque sprint "
        "a suivi le déroulement suivant : planification des tâches, développement, revue de code "
        "avec l'encadrant et rétrospective. L'utilisation de Git pour le contrôle de version et "
        "la CI/CD via GitHub Actions a assuré la traçabilité et la qualité du code tout au long du projet.")

    h(doc, "3.2  Architecture technique détaillée", level=2)
    txt(doc, "L'application repose sur une architecture en couches découplées :")
    doc.add_paragraph()
    txt(doc, "Frontend :", bold=True)
    bul(doc, "Next.js 15 (React 19, TypeScript, Tailwind CSS)")
    bul(doc, "Routage basé sur les rôles (Admin / Utilisateur / Sous-utilisateur)")
    bul(doc, "Services de communication avec l'API via JWT Bearer")
    bul(doc, "Support bilingue français/arabe avec bascule en temps réel")
    bul(doc, "Saisie et restitution vocale via Web Speech API")
    doc.add_paragraph()
    txt(doc, "Backend :", bold=True)
    bul(doc, "FastAPI (Python 3.11) — API RESTful avec documentation automatique")
    bul(doc, "LangGraph — orchestration du pipeline d'analyse à trois agents")
    bul(doc, "DeepSeek Chat — modèle de langage large (via API OpenAI-compatible)")
    bul(doc, "python-docx + Jinja2 — génération de documents Word et rapports HTML")
    bul(doc, "WeasyPrint — conversion HTML vers PDF")
    doc.add_paragraph()
    txt(doc, "Couche de données :", bold=True)
    bul(doc, "PostgreSQL 16 + pgvector — base de données relationnelle avec extension vectorielle")
    bul(doc, "Redis 7 — cache et gestion des sessions")
    doc.add_paragraph()
    txt(doc, "Services IA (locaux) :", bold=True)
    bul(doc, "BAAI/bge-m3 — modèle d'embeddings multilingue (1024 dimensions)")
    bul(doc, "BAAI/bge-reranker-v2-m3 — réordonnancement des résultats de recherche")
    doc.add_paragraph()

    figure(doc, "architecture generale.svg",
           "Figure 1 — Architecture générale de ConformIA", width=Inches(5.8))

    h(doc, "3.3  Spécifications fonctionnelles", level=2)
    txt(doc, "Rôles utilisateurs :", bold=True)
    bul(doc, "Administrateur : accès complet à toutes les fonctionnalités, gestion des utilisateurs et approbation des inscriptions.")
    bul(doc, "Utilisateur : accès à toutes les fonctionnalités d'analyse et de génération, gestion de ses sous-utilisateurs.")
    bul(doc, "Sous-utilisateur : accès aux fonctionnalités d'analyse et de chat, délégué par un utilisateur parent.")
    doc.add_paragraph()
    txt(doc, "Fonctionnalités principales :", bold=True)
    bul(doc, "Analyse de conformité : soumission d'un contrat (PDF/DOCX/TXT), traitement asynchrone, affichage des findings avec verdicts et sévérités.")
    bul(doc, "Génération de rapports : export PDF ou HTML de l'analyse de conformité.")
    bul(doc, "Génération de contrats : création d'un contrat conforme à partir d'une description en langage naturel.")
    bul(doc, "Correction de contrats : application automatique des recommandations sur le contrat original.")
    bul(doc, "Chat juridique : assistant conversationnel général ou contextuel (lié à une analyse).")
    bul(doc, "Gestion des comptes : inscription avec vérification d'email, réinitialisation de mot de passe, modification du profil.")

    h(doc, "3.4  Diagrammes fonctionnels", level=2)

    txt(doc, "3.4.1  Diagramme des cas d'utilisation", bold=True, after=4)
    txt(doc,
        "Le diagramme ci-dessous représente les interactions entre les quatre acteurs du système "
        "(Visiteur, Sous-utilisateur, Utilisateur, Administrateur) et les cas d'utilisation disponibles. "
        "Les relations d'héritage illustrent l'accumulation progressive des droits.",
        after=8)
    figure(doc, "use case diagram.svg",
           "Figure 2 — Diagramme des cas d'utilisation", width=Inches(5.5))

    txt(doc, "3.4.2  Diagramme de classes", bold=True, after=4)
    txt(doc,
        "Le diagramme de classes ci-dessous présente les entités principales du système et leurs "
        "relations. Les classes Utilisateur, Contrat, Analyse, Finding, Citation et ArticleJuridique "
        "constituent le cœur du modèle de données, complétées par les classes TokenEmail et "
        "TokenReinitialisation pour la gestion sécurisée des comptes.",
        after=8)
    figure(doc, "diagram de classe.svg",
           "Figure 3 — Diagramme de classes", width=Inches(5.5))

    txt(doc, "3.4.3  Diagramme de séquence", bold=True, after=4)
    txt(doc,
        "Le diagramme de séquence ci-dessous illustre le flux complet d'une analyse de conformité, "
        "depuis le téléchargement du contrat par l'utilisateur jusqu'à l'affichage des résultats, "
        "en passant par les trois nœuds du pipeline LangGraph et la génération du rapport PDF.",
        after=8)
    figure(doc, "diagrame de sequance.svg",
           "Figure 4 — Diagramme de séquence — Analyse de conformité", width=Inches(5.8))

    h(doc, "3.5  Conclusion", level=2)
    txt(doc,
        "Ce chapitre a décrit l'architecture technique de ConformIA, ses spécifications fonctionnelles "
        "et ses principaux diagrammes de conception. L'architecture découplée (frontend/backend/base "
        "de données) garantit la maintenabilité et l'évolutivité du système. Le pipeline LangGraph "
        "à trois agents, associé au mécanisme de Citation Guard, assure une analyse fiable et "
        "vérifiable de la conformité contractuelle.")

    doc.add_page_break()

    # ──────────────────────────────────────────
    # CHAPITRE IV
    # ──────────────────────────────────────────
    h(doc, "Chapitre IV : Développement et réalisation", level=1)

    h(doc, "4.1  Choix des technologies", level=2)
    txt(doc,
        "Les technologies retenues pour ConformIA ont été sélectionnées sur la base de leur maturité, "
        "de leur performance pour le traitement du langage naturel multilingue, et de leur adéquation "
        "avec les contraintes du projet.")
    doc.add_paragraph()

    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Table Grid'
    th = tech_table.rows[0].cells
    th[0].text = "Couche"
    th[1].text = "Technologie"
    th[2].text = "Justification"
    for cell in th:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    techs = [
        ("Frontend",      "Next.js 15 + React 19",        "SSR, performances, écosystème React"),
        ("Frontend",      "TypeScript + Tailwind CSS",     "Typage fort, design réactif"),
        ("Backend",       "FastAPI (Python 3.11)",         "Performances, async, documentation auto"),
        ("Orchestration", "LangGraph",                     "Pipeline IA stateful et contrôlable"),
        ("LLM",           "DeepSeek Chat",                 "Performances, API OpenAI-compatible, coût"),
        ("Embeddings",    "BAAI/bge-m3",                   "Multilingue, 1024 dims, SOTA"),
        ("Reranker",      "BAAI/bge-reranker-v2-m3",      "Précision top-5, multilingue"),
        ("Base de données","PostgreSQL 16 + pgvector",     "Recherche vectorielle native, fiabilité"),
        ("Cache",         "Redis 7",                       "Légèreté, performances"),
        ("Conteneurs",    "Docker + Docker Compose",       "Portabilité, reproductibilité"),
        ("Doc Word",      "python-docx + Jinja2",          "Génération programmatique de DOCX"),
        ("PDF",           "WeasyPrint",                    "Rendu HTML→PDF fidèle"),
    ]

    for layer, tech, reason in techs:
        row = tech_table.add_row()
        row.cells[0].text = layer
        row.cells[1].text = tech
        row.cells[2].text = reason
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    h(doc, "4.2  Mise en place du backend (FastAPI + LangGraph)", level=2)

    txt(doc, "4.2.1  Structure générale de l'API", bold=True, after=4)
    txt(doc,
        "Le backend est structuré en plusieurs modules distincts : les routes API (analyses, chat, "
        "génération de documents), les routes d'authentification (inscription, connexion, vérification "
        "d'email, réinitialisation de mot de passe), les routes de gestion des utilisateurs, et "
        "les modules partagés (schémas Pydantic, authentification JWT, guards).")

    txt(doc, "4.2.2  Pipeline LangGraph — Agent Extracteur", bold=True, after=4)
    txt(doc,
        "L'agent extracteur utilise DeepSeek Chat avec une sortie structurée (modèle Pydantic "
        "ContratsExtraction) pour identifier les éléments clés du contrat : type (CDI, CDD, CTT, "
        "Stage), parties (employeur, employé), poste, salaire, période d'essai, et liste des clauses "
        "à analyser. Des prompts bilingues avec exemples few-shot garantissent la précision de l'extraction.")

    txt(doc, "4.2.3  Pipeline LangGraph — Agent Récupérateur", bold=True, after=4)
    txt(doc,
        "Pour chaque clause identifiée, l'agent récupérateur génère un vecteur d'embedding (1024 "
        "dimensions) à l'aide du modèle BAAI/bge-m3, puis effectue une recherche par similarité "
        "cosinus dans pgvector sur les 2 176 articles juridiques indexés (top-25). "
        "Un réordonnancement avec BAAI/bge-reranker-v2-m3 sélectionne ensuite les 5 articles les "
        "plus pertinents, qui seront soumis au vérificateur.")

    txt(doc, "4.2.4  Pipeline LangGraph — Agent Vérificateur et Citation Guard", bold=True, after=4)
    txt(doc,
        "L'agent vérificateur soumet chaque clause, accompagnée de ses 5 articles de référence, "
        "au modèle de langage pour obtenir un verdict structuré (CONFORME, NON_CONFORME, EXIGE_REVUE), "
        "une sévérité (BLOQUANT, MAJEUR, MINEUR), une recommandation et des citations précises. "
        "Le Citation Guard valide de manière déterministe que chaque article cité est bien présent "
        "dans le top-5 et que le texte cité correspond exactement au contenu de l'article (normalisation "
        "NFKC + comparaison de sous-chaînes). En cas d'échec, un second essai est effectué avec "
        "un feedback contenant la liste des articles disponibles.")

    txt(doc, "4.2.5  Authentification et sécurité", bold=True, after=4)
    txt(doc,
        "Le système d'authentification repose sur JWT (HS256, expiration 24h) et le hachage bcrypt "
        "des mots de passe. Il intègre un workflow complet de vérification d'adresse email (envoi d'un "
        "lien de confirmation par email) et de réinitialisation de mot de passe (envoi d'un lien "
        "temporaire sécurisé). Les routes API sont protégées par des dépendances FastAPI "
        "(get_current_user, require_admin, require_approved) garantissant un contrôle d'accès strict.")

    h(doc, "4.3  Développement du frontend (Next.js 15)", level=2)
    txt(doc,
        "L'interface utilisateur est développée avec Next.js 15 et React 19 en TypeScript, stylisée "
        "avec Tailwind CSS. Elle adopte une architecture de composants réutilisables et un contexte "
        "d'authentification (AuthProvider) gérant l'état de session via localStorage.")
    doc.add_paragraph()
    txt(doc,
        "La page principale (page.tsx) constitue le cœur de l'application : elle offre une interface "
        "de chat intégrée permettant à la fois les questions juridiques générales, l'envoi d'un contrat "
        "pour analyse, et les interactions contextuelles sur les résultats. Elle inclut également la "
        "détection automatique des intentions (génération ou correction de contrat) via des "
        "correspondances de mots-clés.")
    doc.add_paragraph()
    txt(doc,
        "Le support bilingue (français/arabe) est implémenté via une bascule de langue en en-tête, "
        "qui adapte dynamiquement les prompts envoyés à l'API et la direction du texte (RTL pour l'arabe). "
        "La saisie et la restitution vocales utilisent l'API Web Speech (SpeechRecognition et "
        "SpeechSynthesis) avec les locales fr-FR et ar-SA.")

    h(doc, "4.4  Corpus juridique et système RAG", level=2)
    txt(doc,
        "Le corpus juridique a été constitué à partir des textes officiels mauritaniens numérisés, "
        "puis segmenté en 2 176 articles individuels grâce à un parser basé sur des expressions "
        "régulières hiérarchiques (Livre → Titre → Chapitre → Section → Article). Chaque article "
        "est enrichi de métadonnées (juridiction, code, numéro d'article, chemin hiérarchique) "
        "avant son indexation.")
    doc.add_paragraph()
    txt(doc,
        "Les embeddings sont générés en lot (batch_size=32) avec normalisation L2, garantissant "
        "des comparaisons cosinus précises. L'ingestion est conçue pour être reprise après "
        "interruption (idempotente via comptage des articles existants) et traite les articles "
        "par chunks de 200 pour éviter les débordements mémoire.")

    h(doc, "4.5  Tests et intégration continue", level=2)
    txt(doc,
        "Le projet inclut une suite de 13 modules de tests couvrant les schémas Pydantic, "
        "le Citation Guard, les agents individuels (extracteur, récupérateur, vérificateur), "
        "le pipeline complet, les routes API, la génération de rapports et les opérations de base "
        "de données. Les tests de schémas et du Citation Guard sont exécutés automatiquement "
        "à chaque commit via GitHub Actions (Python 3.11, ruff pour le linting).")

    h(doc, "4.6  Présentation de l'interface utilisateur", level=2)

    txt(doc, "4.6.1  Authentification et gestion du compte", bold=True, after=4)
    txt(doc,
        "Les figures suivantes illustrent les différentes étapes du flux d'authentification : "
        "connexion, inscription, vérification d'email et réinitialisation du mot de passe.",
        after=8)

    figure(doc, "Screenshot 2026-06-04 235949.png",
           "Figure 5 — Page de connexion", width=Inches(5.5))
    figure(doc, "Screenshot 2026-06-04 235958.png",
           "Figure 6 — Page d'inscription", width=Inches(5.5))
    figure(doc, "Screenshot 2026-06-06 134022.png",
           "Figure 7 — Page de vérification d'adresse email", width=Inches(5.0))
    figure(doc, "Screenshot 2026-06-06 133836.png",
           "Figure 8 — Page « Mot de passe oublié »", width=Inches(5.0))
    figure(doc, "Screenshot 2026-06-06 134437.png",
           "Figure 9 — Page de réinitialisation du mot de passe", width=Inches(5.0))

    txt(doc, "4.6.2  Interface principale — Assistant juridique", bold=True, after=4)
    txt(doc,
        "L'interface principale de ConformIA présente un assistant juridique conversationnel. "
        "L'utilisateur peut poser des questions sur le droit du travail mauritanien, télécharger "
        "un contrat pour analyse, ou demander la génération d'un nouveau contrat. La barre latérale "
        "affiche l'historique des analyses en cours ou terminées.",
        after=8)
    figure(doc, "Screenshot 2026-06-05 201954.png",
           "Figure 10 — Interface principale de l'assistant juridique ConformIA", width=Inches(5.8))

    txt(doc, "4.6.3  Résultats d'analyse et génération de documents", bold=True, after=4)
    txt(doc,
        "Après traitement du contrat par le pipeline LangGraph, les résultats s'affichent sous "
        "forme d'un tableau de conformité indiquant pour chaque clause : le verdict "
        "(EXIGE_REVUE, CONFORME, NON_CONFORME), la sévérité (BLOQUANT, MAJEUR, MINEUR) et "
        "l'article juridique de référence. L'utilisateur peut ensuite télécharger le rapport PDF "
        "ou demander la correction automatique du contrat.",
        after=8)
    figure(doc, "Screenshot 2026-06-05 202308.png",
           "Figure 11 — Résultats d'analyse de conformité contractuelle", width=Inches(5.8))
    figure(doc, "Screenshot 2026-06-05 202559.png",
           "Figure 12 — Génération et téléchargement du contrat corrigé", width=Inches(5.8))

    txt(doc, "4.6.4  Tableau de bord administrateur", bold=True, after=4)
    txt(doc,
        "Le tableau de bord administrateur offre une vue synthétique de la plateforme : nombre "
        "d'utilisateurs en attente d'approbation, utilisateurs validés et sous-utilisateurs. "
        "L'administrateur peut approuver ou rejeter les inscriptions directement depuis l'interface.",
        after=8)
    figure(doc, "Screenshot 2026-06-05 201748.png",
           "Figure 13 — Tableau de bord administrateur", width=Inches(5.8))
    figure(doc, "Screenshot 2026-06-05 201819.png",
           "Figure 14 — Gestion des approbations en attente", width=Inches(5.8))

    txt(doc, "4.6.5  Gestion des sous-utilisateurs et paramètres", bold=True, after=4)
    txt(doc,
        "Un utilisateur peut créer et gérer des sous-comptes (sous-utilisateurs) depuis la section "
        "dédiée. La page Paramètres permet la modification du profil (nom, email) et le changement "
        "du mot de passe.",
        after=8)
    figure(doc, "Screenshot 2026-06-05 202008.png",
           "Figure 15 — Gestion des sous-utilisateurs", width=Inches(5.8))
    figure(doc, "Screenshot 2026-06-05 202029.png",
           "Figure 16 — Paramètres du compte utilisateur", width=Inches(5.8))

    h(doc, "4.7  Conclusion", level=2)
    txt(doc,
        "Ce chapitre a présenté en détail les choix technologiques de ConformIA, l'implémentation "
        "du backend (FastAPI, pipeline LangGraph, Citation Guard, authentification sécurisée), "
        "le développement du frontend bilingue et l'ensemble des interfaces réalisées. "
        "La combinaison de technologies modernes d'IA (LLM + RAG + reranking) avec une architecture "
        "web robuste a permis de livrer une application fonctionnelle, fiable et ergonomique, "
        "répondant pleinement aux objectifs fixés en début de projet.")

    doc.add_page_break()

    # ──────────────────────────────────────────
    # CONCLUSION GENERALE
    # ──────────────────────────────────────────
    h(doc, "Conclusion générale", level=1)
    doc.add_paragraph()

    txt(doc, "Résultats obtenus", bold=True, size=13, after=6)
    txt(doc,
        "À l'issue de ce projet de fin d'études, nous avons livré une application web complète et "
        "fonctionnelle répondant à l'ensemble des objectifs initialement définis. Parmi les "
        "réalisations notables :")
    bul(doc, "Un pipeline d'analyse IA à trois agents (extracteur, récupérateur, vérificateur) orchestré par LangGraph, capable d'analyser la conformité de tout contrat de travail mauritanien en quelques minutes.")
    bul(doc, "Un corpus juridique de 2 176 articles indexés et recherchables sémantiquement via pgvector, couvrant les cinq principaux textes législatifs mauritaniens.")
    bul(doc, "Un mécanisme de Citation Guard garantissant que toutes les citations légales produites par le système sont vérifiables et exactes, éliminant ainsi le risque d'hallucinations.")
    bul(doc, "Une interface utilisateur bilingue (français/arabe) avec saisie vocale, accessible et intuitive pour des utilisateurs non techniciens.")
    bul(doc, "Un système de gestion des utilisateurs à trois niveaux (Admin, Utilisateur, Sous-utilisateur) avec workflow d'approbation, vérification d'email et réinitialisation de mot de passe.")
    bul(doc, "Des fonctionnalités de génération et de correction automatique de contrats au format DOCX, ainsi qu'un export de rapports de conformité en PDF.")
    bul(doc, "Un déploiement conteneurisé via Docker Compose garantissant la portabilité et la reproductibilité de l'environnement.")
    doc.add_paragraph()

    txt(doc, "Difficultés rencontrées", bold=True, size=13, after=6)
    txt(doc,
        "La réalisation de ce projet n'a pas été sans défis techniques et organisationnels :")
    bul(doc, "Gestion des hallucinations du LLM : les modèles de langage tendent à inventer des citations légales plausibles mais inexactes. La conception du Citation Guard a nécessité une réflexion approfondie sur les mécanismes de validation déterministe.")
    bul(doc, "Performance du pipeline d'embeddings : le traitement de 2 176 articles avec le modèle BGE-M3 a nécessité une gestion fine de la mémoire (chunking par 200 articles) et de la reprise après interruption.")
    bul(doc, "Support multilingue (arabe) : l'intégration de l'arabe dans l'interface, notamment pour la saisie et la restitution vocales et la direction RTL du texte, a représenté un défi d'ergonomie et de compatibilité.")
    bul(doc, "Coordination d'équipe : la répartition efficace du travail entre les membres de l'équipe sur un projet aussi vaste a nécessité une organisation rigoureuse et une communication constante.")
    doc.add_paragraph()

    txt(doc, "Solutions apportées", bold=True, size=13, after=6)
    bul(doc, "Citation Guard déterministe : validation par normalisation NFKC et comparaison exacte de sous-chaînes, avec mécanisme de retry (une tentative) et fallback vers le verdict EXIGE_REVUE.")
    bul(doc, "Ingestion par chunks avec reprise idempotente : traitement des articles par lots de 200 avec comptage des articles existants pour permettre la reprise sans duplication.")
    bul(doc, "Localisation dynamique : système de bascule fr-FR/ar-SA adaptatif, affectant à la fois l'interface, les prompts envoyés au LLM et les paramètres de reconnaissance vocale.")
    bul(doc, "Méthodologie Scrum : les sprints bihebdomadaires avec revue d'encadrant ont permis d'identifier et de résoudre rapidement les blocages techniques.")
    doc.add_paragraph()

    txt(doc, "Perspectives d'amélioration", bold=True, size=13, after=6)
    txt(doc,
        "Ce projet ouvre plusieurs perspectives d'évolution intéressantes pour les versions futures "
        "de ConformIA :")
    bul(doc, "Extension du corpus à d'autres pays d'Afrique de l'Ouest (Sénégal, Mali, Côte d'Ivoire) pour créer une plateforme régionale de conformité contractuelle.")
    bul(doc, "Intégration d'un module de statistiques avancées et d'analytique permettant aux entreprises de suivre leur niveau de conformité dans le temps.")
    bul(doc, "Développement d'une application mobile (React Native) pour l'accès à l'assistant juridique en déplacement.")
    bul(doc, "Fine-tuning d'un modèle de langage spécialisé sur le droit mauritanien pour améliorer encore la précision des analyses.")
    bul(doc, "Intégration avec des systèmes de gestion RH existants (API) pour automatiser la vérification à chaque création ou modification de contrat.")
    bul(doc, "Ajout d'un module de veille juridique alertant les utilisateurs lors de modifications législatives affectant leurs contrats existants.")
    doc.add_paragraph()

    txt(doc,
        "ConformIA représente une contribution concrète à la modernisation de la pratique juridique "
        "en Mauritanie et démontre la faisabilité et la pertinence de l'application des technologies "
        "d'intelligence artificielle au service du droit. Ce projet constitue pour notre équipe "
        "une expérience formatrice exceptionnelle, alliant rigueur technique, créativité et sens "
        "des responsabilités sociales.",
        italic=True)

    doc.add_page_break()

    # ──────────────────────────────────────────
    # BIBLIOGRAPHIE
    # ──────────────────────────────────────────
    h(doc, "Bibliographie", level=1)
    doc.add_paragraph()

    refs = [
        "[1]  Mauritanie, Loi N° 2004-017 portant Code du Travail de la République Islamique de Mauritanie, Journal Officiel, 2004.",
        "[2]  Mauritanie, Ordonnance N° 89-126 portant Code des Obligations et des Contrats, modifiée par la Loi N° 2001-31, 2001.",
        "[3]  Mauritanie, Loi N° 2000-05 portant Code du Commerce, 2000.",
        "[4]  UNICEMA / UTM, Convention Collective Générale du Travail, Mauritanie.",
        "[5]  Organisation Internationale du Travail (OIT), Conventions Internationales ratifiées par la Mauritanie, www.ilo.org.",
        "[6]  LangChain, LangGraph — Build reliable, stateful AI agents, Documentation officielle, 2024. [En ligne] https://langchain-ai.github.io/langgraph/",
        "[7]  FastAPI, FastAPI — Modern, fast web framework for building APIs with Python, Documentation officielle, 2024. [En ligne] https://fastapi.tiangolo.com/",
        "[8]  Next.js, The React Framework for the Web, Documentation officielle, Vercel, 2024. [En ligne] https://nextjs.org/docs",
        "[9]  BAAI, BGE M3 — Multi-Functionality, Multi-Linguality, Multi-Granularity Text Embeddings Through Self-Knowledge Distillation, arXiv:2309.07597, 2023.",
        "[10] pgvector, Open-source vector similarity search for Postgres, GitHub, 2024. [En ligne] https://github.com/pgvector/pgvector",
        "[11] DeepSeek, DeepSeek-V2: A Strong, Economical, and Efficient Mixture-of-Experts Language Model, arXiv:2405.04434, 2024.",
        "[12] Lewis, P. et al., Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks, NeurIPS, 2020.",
        "[13] python-docx, Python library for creating and updating Microsoft Word (.docx) files, Documentation officielle. [En ligne] https://python-docx.readthedocs.io/",
    ]

    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.paragraph_format.left_indent = Cm(0.8)
        p_ref.paragraph_format.first_line_indent = Cm(-0.8)
        r = p_ref.add_run(ref)
        r.font.size = Pt(11)

    # ──────────────────────────────────────────
    # SAVE
    # ──────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"OK - Rapport genere : {OUTPUT}")

if __name__ == "__main__":
    build()
