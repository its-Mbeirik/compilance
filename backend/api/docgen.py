"""
Génération de documents Word (.docx) pour contrats mauritaniens.
  - generate_contract_docx(description) → (title, bytes)
  - correct_contract_docx(analysis_rec, original_text) → (title, bytes)
"""
import io
import os
import re
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Low-level XML helpers ────────────────────────────────────────────────────

def _rfonts(run, name="Calibri"):
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rPr.insert(0, rFonts)


def _spacing(p, before_pt=None, after_pt=None, line_multiple=None):
    """Set paragraph spacing via XML (reliable across python-docx versions)."""
    pPr = p._p.get_or_add_pPr()
    old = pPr.find(qn("w:spacing"))
    if old is not None:
        pPr.remove(old)
    el = OxmlElement("w:spacing")
    if before_pt is not None:
        el.set(qn("w:before"), str(round(before_pt * 20)))
    if after_pt is not None:
        el.set(qn("w:after"), str(round(after_pt * 20)))
    if line_multiple is not None:
        el.set(qn("w:line"), str(round(line_multiple * 240)))
        el.set(qn("w:lineRule"), "auto")
    pPr.append(el)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)


# ── Paragraph builders ───────────────────────────────────────────────────────

def _run(p, text, size=11, bold=False, color=None, font="Calibri"):
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    _rfonts(run, font)
    return run


def _add_cover_page(doc: Document, title: str, employer: str, employee: str, date_str: str):
    """Full cover page followed by a page break."""
    # Push content down with top spacing
    p_top = doc.add_paragraph()
    _spacing(p_top, before_pt=100, after_pt=0)

    # Document type label
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_label, "CONTRAT DE TRAVAIL", size=12, bold=False, color=RGBColor(0x60, 0x60, 0x60))
    _spacing(p_label, before_pt=0, after_pt=8)

    # Main title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_title, title.upper(), size=20, bold=True)
    _spacing(p_title, before_pt=0, after_pt=6)

    # Subtitle line
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_sub, "Version corrigée — Conforme au droit mauritanien", size=11, color=RGBColor(0x60, 0x60, 0x60))
    _spacing(p_sub, before_pt=0, after_pt=40)

    # Horizontal line via bottom border on paragraph
    p_line = doc.add_paragraph()
    pPr = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    _spacing(p_line, before_pt=0, after_pt=30)

    # Info block
    for label, value in [
        ("Employeur", employer or "—"),
        ("Salarié",   employee or "—"),
        ("Date",      date_str or date.today().strftime("%d/%m/%Y")),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"{label} : ", size=12, bold=True)
        _run(p, value, size=12)
        _spacing(p, before_pt=0, after_pt=8)

    doc.add_page_break()


def _add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, text.upper(), size=16, bold=True)
    _spacing(p, before_pt=6, after_pt=24)


def _add_section_heading(doc: Document, text: str):
    """## Section — bold 12pt, before 12pt, after 6pt."""
    p = doc.add_paragraph()
    _run(p, text, size=12, bold=True)
    _spacing(p, before_pt=12, after_pt=6)


def _add_sub_heading(doc: Document, text: str):
    """### Sub-heading — bold 11pt, before 8pt, after 4pt."""
    p = doc.add_paragraph()
    _run(p, text, size=11, bold=True)
    _spacing(p, before_pt=8, after_pt=4)


def _add_article_heading(doc: Document, number: str, title: str):
    """Article N — Title: bold 12pt, before 14pt, after 6pt."""
    p = doc.add_paragraph()
    label = f"Article {number}"
    if title.strip():
        label += f" — {title.strip()}"
    _run(p, label, size=12, bold=True)
    _spacing(p, before_pt=14, after_pt=6, line_multiple=1.15)


def _add_body(doc: Document, text: str):
    """Body paragraph: 11pt, justified, 1.15 line spacing."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, size=11)
    _spacing(p, before_pt=0, after_pt=4, line_multiple=1.15)


def _add_bullet(doc: Document, text: str):
    """Bullet item: 11pt, indented, 4pt after."""
    p = doc.add_paragraph()
    _run(p, f"• {text}", size=11)
    pf = p.paragraph_format
    pf.left_indent = Cm(1.0)
    _spacing(p, before_pt=0, after_pt=4, line_multiple=1.15)


def _add_signature_block(doc: Document, employer: str, employee: str):
    """Two-column signature block at the end of the document."""
    # Spacer
    p_sp = doc.add_paragraph()
    _spacing(p_sp, before_pt=30, after_pt=0)

    # "Signatures" heading
    p_h = doc.add_paragraph()
    _run(p_h, "SIGNATURES", size=12, bold=True)
    _spacing(p_h, before_pt=0, after_pt=14)

    table = doc.add_table(rows=4, cols=2)
    _remove_table_borders(table)

    # Set equal column widths
    for cell in table.columns[0].cells:
        cell.width = Cm(8)
    for cell in table.columns[1].cells:
        cell.width = Cm(8)

    rows_data = [
        (f"Pour l'Employeur :", f"Le Salarié :", True),
        (employer or "_______________", employee or "_______________", False),
        ("", "", False),
        ("Date : ___/___/______\n\nSignature :\n\n\n_______________",
         "Date : ___/___/______\n\nSignature :\n\n\n_______________", False),
    ]

    for i, (left_text, right_text, bold) in enumerate(rows_data):
        left_cell  = table.rows[i].cells[0]
        right_cell = table.rows[i].cells[1]
        for cell, text in [(left_cell, left_text), (right_cell, right_text)]:
            cp = cell.paragraphs[0]
            _run(cp, text, size=11, bold=bold)
            _spacing(cp, before_pt=0, after_pt=6)


# ── Document builder ─────────────────────────────────────────────────────────

def _build_docx(
    title: str,
    llm_text: str,
    employer: str = "",
    employee: str = "",
    date_str: str = "",
    cover_page: bool = False,
) -> bytes:
    """
    Parse LLM output (Markdown-lite) and build a professionally formatted .docx.

    Markers:
      ## Section title       → bold section heading
      ### Sub-heading        → smaller bold heading
      Article N — Title      → article heading (bold 12pt) + body on next lines
      - bullet / • bullet    → indented bullet
      plain text             → justified body paragraph
    """
    doc = Document()

    # 2.5 cm margins on all sides
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    if cover_page:
        _add_cover_page(doc, title, employer, employee, date_str)

    _add_title(doc, title)

    _ARTICLE_RE = re.compile(r"^Art(?:icle)?\.?\s+(\d+)\s*[:\-–—]\s*(.*)", re.IGNORECASE)
    _ARTICLE_NUM_RE = re.compile(r"^Art(?:icle)?\.?\s+(\d+)\s*:?\s*(.*)", re.IGNORECASE)

    for line in llm_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue  # skip blank lines (spacing is handled per-paragraph)

        if stripped.startswith("## "):
            _add_section_heading(doc, stripped[3:])

        elif stripped.startswith("### "):
            _add_sub_heading(doc, stripped[4:])

        elif _ARTICLE_RE.match(stripped):
            m = _ARTICLE_RE.match(stripped)
            _add_article_heading(doc, m.group(1), m.group(2))

        elif _ARTICLE_NUM_RE.match(stripped):
            m = _ARTICLE_NUM_RE.match(stripped)
            _add_article_heading(doc, m.group(1), m.group(2))

        elif stripped.startswith(("- ", "• ", "* ")):
            _add_bullet(doc, stripped[2:])

        else:
            _add_body(doc, stripped)

    _add_signature_block(doc, employer, employee)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── LLM client ───────────────────────────────────────────────────────────────

def _llm(temperature=0.2, max_tokens=8192):
    from langchain_openai import ChatOpenAI
    return ChatOpenAI(
        model=os.getenv("LLM_MODEL", "deepseek-chat"),
        api_key=os.getenv("DEEPSEEK_API_KEY") or os.getenv("OPENAI_API_KEY"),
        base_url=os.getenv("OPENAI_API_BASE", "https://api.deepseek.com"),
        temperature=temperature,
        max_tokens=max_tokens,
    )


# ── Public API ────────────────────────────────────────────────────────────────

_PREVIEW_SYSTEM = {
    "fr": (
        "Tu es un expert juridique spécialisé en droit mauritanien "
        "(Code du Travail Loi N° 2004-017, COC, Convention Collective Générale du Travail).\n\n"
        "Pour chaque clause non conforme, génère une proposition de correction "
        "avec ce format markdown exact — sans introduction, sans conclusion :\n\n"
        "**📌 Article [N] — [Nom de la clause]**\n"
        "❌ **Problème détecté :** [explication courte et précise]\n"
        "✅ **Clause corrigée proposée :**\n"
        "*\"[Texte complet et juridiquement conforme de la clause corrigée]\"*\n"
        "📖 **Référence légale :** Art. [N] du [Code source]\n\n"
        "---\n\n"
        "Règles absolues :\n"
        "- La devise est TOUJOURS MRU (Ouguiya mauritanien). Jamais FCFA ni XOF.\n"
        "- Traiter UNIQUEMENT les clauses NON_CONFORME ou EXIGE_REVUE.\n"
        "- Ne pas répéter les clauses CONFORME.\n"
        "- Être concis, précis et pratique dans la correction proposée."
    ),
    "ar": (
        "أنت خبير قانوني متخصص في القانون الموريتاني "
        "(قانون العمل رقم 2004-017، ومدونة الالتزامات والعقود، والاتفاقية الجماعية العامة للعمل).\n\n"
        "لكل بند غير مطابق، أنشئ اقتراح تصحيح بهذا التنسيق الحرفي:\n\n"
        "**📌 المادة [N] — [اسم البند]**\n"
        "❌ **المشكلة المكتشفة:** [شرح مختصر ودقيق]\n"
        "✅ **البند المصحَّح المقترح:**\n"
        "*\"[النص الكامل والمطابق قانونياً للبند المصحَّح]\"*\n"
        "📖 **المرجع القانوني:** المادة [N] من [المصدر]\n\n"
        "---\n\n"
        "قواعد مطلقة:\n"
        "- العملة دائماً MRU (أوقية موريتانية). لا FCFA.\n"
        "- تناوَل فقط البنود NON_CONFORME أو EXIGE_REVUE.\n"
        "- لا تكرر البنود المطابقة."
    ),
}


def generate_correction_preview_text(
    analysis_rec: dict,
    original_text: str,
    language: str = "fr",
) -> str:
    """
    Generate a markdown-formatted chat message showing corrected clauses.
    Ends with the confirmation question. Does NOT produce a .docx file.
    """
    from langchain_core.messages import HumanMessage, SystemMessage
    lang = language if language in _PREVIEW_SYSTEM else "fr"

    findings  = analysis_rec.get("findings", [])
    extracted = analysis_rec.get("extracted", {})

    non_conformes = [f for f in findings if f.get("verdict") in ("NON_CONFORME", "EXIGE_REVUE")]

    if not non_conformes:
        if lang == "ar":
            return "✅ جميع البنود المحللة مطابقة للقانون الموريتاني. لا يلزم إجراء أي تصحيح."
        return "✅ Toutes les clauses analysées sont conformes au droit mauritanien. Aucune correction n'est nécessaire."

    salary = extracted.get("salaire_mensuel_mru") or extracted.get("salaire_mensuel_fcfa")
    salary_str = f"{salary:,.0f} MRU" if salary else "non précisé"

    findings_for_prompt = "\n".join(
        f"- Clause {f.get('clause_id')} ({f.get('verdict')}): "
        f"{f.get('recommendation', '')} [Réf: {f.get('cited_article_id', '')}]"
        for f in non_conformes
    )

    if lang == "ar":
        prompt = (
            f"مقتطف من العقد الأصلي:\n{original_text[:4000]}\n\n"
            f"البيانات المستخرجة:\n"
            f"- نوع العقد: {extracted.get('type_contrat','?')}\n"
            f"- جهة العمل: {extracted.get('employeur','?')}\n"
            f"- الموظف: {extracted.get('employe','?')}\n"
            f"- الراتب الشهري: {salary_str}\n\n"
            f"البنود غير المطابقة:\n{findings_for_prompt}\n\n"
            "أنشئ اقتراحات التصحيح بالتنسيق المطلوب."
        )
        confirm_q = "\n\n---\n\nهل تريد مني إنشاء مستند .docx المصحَّح بهذه التعديلات؟"
    else:
        prompt = (
            f"Extrait du contrat original :\n{original_text[:4000]}\n\n"
            f"Données extraites :\n"
            f"- Type de contrat : {extracted.get('type_contrat','?')}\n"
            f"- Employeur : {extracted.get('employeur','?')}\n"
            f"- Salarié : {extracted.get('employe','?')}\n"
            f"- Salaire mensuel brut : {salary_str}\n\n"
            f"Clauses non conformes ({len(non_conformes)}) :\n{findings_for_prompt}\n\n"
            "Génère les propositions de correction au format demandé."
        )
        confirm_q = "\n\n---\n\nSouhaitez-vous que je génère le document .docx corrigé avec ces modifications ?"

    resp = _llm(temperature=0.2, max_tokens=3000).invoke([
        SystemMessage(content=_PREVIEW_SYSTEM[lang]),
        HumanMessage(content=prompt),
    ])
    return resp.content.strip() + confirm_q


_GENERATE_SYSTEM = {
    "fr": (
        "Tu es un expert juridique spécialisé en droit mauritanien (Code du Travail Loi N° 2004-017, "
        "Code des Obligations et des Contrats, Convention Collective Générale du Travail). "
        "Génère un contrat complet, structuré et conforme à la législation mauritanienne en vigueur. "
        "Structure le document avec des sections claires marquées ## pour les titres principaux "
        "et ### pour les sous-titres. Numérote chaque clause : 'Article N — Titre du clauses' suivi "
        "du texte de la clause sur les lignes suivantes. "
        "La devise est toujours MRU (Ouguiya mauritanien). "
        "Inclus toutes les mentions légales obligatoires. Réponds uniquement avec le texte du contrat."
    ),
    "ar": (
        "أنت خبير قانوني متخصص في القانون الموريتاني (قانون العمل رقم 2004-017، "
        "ومدونة الالتزامات والعقود، والاتفاقية الجماعية العامة للعمل). "
        "أنشئ عقداً كاملاً ومنظَّماً ومتوافقاً مع التشريعات الموريتانية النافذة. "
        "نظِّم الوثيقة بأقسام واضحة مُعلَّمة بـ ## للعناوين الرئيسية و### للعناوين الفرعية. "
        "رقِّم كل بند: 'المادة N — عنوان البند' ثم نص البند في الأسطر التالية. "
        "العملة دائماً MRU (أوقية موريتانية). "
        "أدرج جميع البيانات القانونية الإلزامية. أجب بنص العقد فقط."
    ),
}

_CORRECT_SYSTEM = {
    "fr": (
        "Tu es un expert juridique spécialisé en droit mauritanien "
        "(Code du Travail Loi N° 2004-017, COC, Convention Collective Générale du Travail).\n\n"
        "Ta mission : produire la version CORRIGÉE et COMPLÈTE d'un contrat de travail mauritanien.\n\n"
        "RÈGLES ABSOLUES — respecte-les sans exception :\n"
        "1. Prends le CONTRAT ORIGINAL comme base unique. Reproduis INTÉGRALEMENT tous ses articles.\n"
        "2. Pour chaque clause marquée CONFORME → copie le texte original tel quel, sans modification.\n"
        "3. Pour chaque clause marquée NON_CONFORME ou EXIGE_REVUE → applique la recommandation fournie "
        "tout en conservant la structure et la numérotation de l'article.\n"
        "4. Remplace chaque [...] par la valeur correspondante extraite (employeur, salarié, salaire, dates…).\n"
        "5. La devise est TOUJOURS en MRU (Ouguiya mauritanien). N'utilise jamais FCFA ni XOF.\n"
        "6. Conserve EXACTEMENT le type de contrat (CDI reste CDI, CDD reste CDD).\n"
        "7. Format : '## Section' pour les sections, 'Article N — Titre' puis corps sur lignes suivantes.\n"
        "8. Tous les articles du contrat original doivent être présents et complets dans le résultat.\n"
        "9. Fournis UNIQUEMENT le texte du contrat corrigé. Aucun commentaire, aucune explication."
    ),
    "ar": (
        "أنت خبير قانوني متخصص في القانون الموريتاني "
        "(قانون العمل رقم 2004-017، ومدونة الالتزامات والعقود، والاتفاقية الجماعية العامة للعمل).\n\n"
        "مهمتك: إنتاج النسخة المصحَّحة والمتكاملة من عقد العمل الموريتاني.\n\n"
        "قواعد مطلقة:\n"
        "1. اتخذ العقد الأصلي قاعدةً وحيدة. استنسخ جميع مواده بالكامل.\n"
        "2. كل بند مُصنَّف CONFORME → انسخ النص الأصلي كما هو دون أي تعديل.\n"
        "3. كل بند NON_CONFORME أو EXIGE_REVUE → طبِّق التوصية المقدَّمة مع الحفاظ على بنية المادة وترقيمها.\n"
        "4. استبدل كل [...] بالقيمة المستخرجة المقابلة.\n"
        "5. العملة دائماً MRU (أوقية موريتانية). لا تستخدم FCFA أبداً.\n"
        "6. احتفظ تماماً بنوع العقد الأصلي.\n"
        "7. التنسيق: '## قسم' للأقسام، 'المادة N — العنوان' ثم النص في الأسطر التالية.\n"
        "8. يجب أن تكون جميع مواد العقد الأصلي حاضرة ومكتملة في النتيجة.\n"
        "9. أجب بنص العقد المصحَّح فقط. لا تعليقات ولا شروحات."
    ),
}


def generate_contract_docx(description: str, language: str = "fr") -> tuple[str, bytes]:
    """Generate a new complete contract from a description. Returns (title, docx_bytes)."""
    from langchain_core.messages import HumanMessage, SystemMessage
    lang = language if language in _GENERATE_SYSTEM else "fr"
    prefix = "أنشئ العقد التالي" if lang == "ar" else "Génère le contrat suivant"

    resp = _llm().invoke([
        SystemMessage(content=_GENERATE_SYSTEM[lang]),
        HumanMessage(content=f"{prefix} : {description}"),
    ])
    content = resp.content

    first_line = next((l.strip().lstrip("#").strip() for l in content.splitlines() if l.strip()), description)
    title = first_line if len(first_line) < 80 else description[:80]
    return title, _build_docx(title, content)


def correct_contract_docx(
    analysis_rec: dict,
    original_text: str,
    language: str = "fr",
) -> tuple[str, bytes]:
    """
    Generate a corrected version of a contract by applying findings recommendations.
    original_text: full text of the uploaded contract (used as base by the LLM).
    Returns (title, docx_bytes).
    """
    from langchain_core.messages import HumanMessage, SystemMessage
    lang = language if language in _CORRECT_SYSTEM else "fr"

    findings  = analysis_rec.get("findings", [])
    extracted = analysis_rec.get("extracted", {})

    employer = extracted.get("employeur", "")
    employee = extracted.get("employe", "")
    salary   = extracted.get("salaire_mensuel_mru") or extracted.get("salaire_mensuel_fcfa")
    salary_str = f"{salary:,.0f} MRU" if salary else "non précisé"

    # Classify findings by verdict
    non_conformes = [f for f in findings if f.get("verdict") in ("NON_CONFORME", "EXIGE_REVUE")]
    conformes     = [f for f in findings if f.get("verdict") == "CONFORME"]

    corrections_text = "\n".join(
        f"  • Clause {f.get('clause_id')} ({f.get('verdict')}) → {f.get('recommendation', 'À corriger')}"
        for f in non_conformes
    ) or "  Aucune correction requise."

    conformes_text = "\n".join(
        f"  • Clause {f.get('clause_id')} → à conserver tel quel"
        for f in conformes
    ) or ""

    if lang == "ar":
        prompt = (
            f"العقد الأصلي :\n{original_text}\n\n"
            f"البيانات المستخرجة :\n"
            f"- نوع العقد : {extracted.get('type_contrat', '?')}\n"
            f"- جهة العمل : {employer}\n"
            f"- الموظف : {employee}\n"
            f"- المنصب : {extracted.get('poste', '?')}\n"
            f"- تاريخ البدء : {extracted.get('date_debut', '?')}\n"
            f"- الراتب الشهري : {salary_str}\n"
            f"- فترة التجربة : {extracted.get('periode_essai_mois', '?')} أشهر\n\n"
            f"التصحيحات الواجب تطبيقها :\n{corrections_text}\n\n"
            f"البنود المطابقة (يجب نسخها كما هي) :\n{conformes_text}\n\n"
            "أنتج الآن النسخة المصحَّحة الكاملة من العقد مع جميع موادّه."
        )
        default_title = "عقد مُصحَّح"
    else:
        prompt = (
            f"CONTRAT ORIGINAL :\n{original_text}\n\n"
            f"DONNÉES EXTRAITES :\n"
            f"- Type de contrat : {extracted.get('type_contrat', '?')}\n"
            f"- Employeur : {employer}\n"
            f"- Salarié : {employee}\n"
            f"- Poste : {extracted.get('poste', '?')}\n"
            f"- Date de début : {extracted.get('date_debut', '?')}\n"
            f"- Salaire mensuel brut : {salary_str}\n"
            f"- Période d'essai : {extracted.get('periode_essai_mois', '?')} mois\n\n"
            f"CORRECTIONS À APPLIQUER (clauses NON_CONFORME / EXIGE_REVUE) :\n{corrections_text}\n\n"
            f"CLAUSES CONFORMES (à recopier sans modification) :\n{conformes_text}\n\n"
            "Produis maintenant le contrat corrigé COMPLET avec tous ses articles."
        )
        default_title = "Contrat corrigé"

    resp = _llm(temperature=0.15, max_tokens=8192).invoke([
        SystemMessage(content=_CORRECT_SYSTEM[lang]),
        HumanMessage(content=prompt),
    ])
    content = resp.content

    # Derive title from type + employee name
    contract_type = extracted.get("type_contrat", "CDI")
    if employee:
        title = f"Contrat {contract_type} — {employee}"
    else:
        first_line = next((l.strip().lstrip("#").strip() for l in content.splitlines() if l.strip()), default_title)
        title = first_line[:80] if len(first_line) < 80 else default_title

    date_str = extracted.get("date_debut") or date.today().strftime("%d/%m/%Y")
    return title, _build_docx(
        title,
        content,
        employer=employer,
        employee=employee,
        date_str=date_str,
        cover_page=True,
    )
