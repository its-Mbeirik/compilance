"""Convert Markdown text to a python-docx Document.

Supports: # / ## / ### / #### headings, bullet lists (- / *), numbered lists
(1. 2.), bold (**...**), italic (*...*), inline code (`...`), horizontal rules
(---), and blank lines. Good enough for compliance reports and corrected
contracts — not a full CommonMark renderer.
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor

_HEADING_RE = re.compile(r"^(#{1,4})\s+(.*)$")
_BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")
_NUMBERED_RE = re.compile(r"^\s*(\d+)\.\s+(.*)$")
_HR_RE = re.compile(r"^\s*---+\s*$")

# Inline emphasis (handled in this order: bold > italic > code).
_INLINE_RE = re.compile(
    r"(\*\*[^*\n]+\*\*)|(\*[^*\n]+\*)|(`[^`\n]+`)"
)


def _add_runs(paragraph, text: str) -> None:
    """Add runs to `paragraph`, applying bold/italic/code styling inline."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def markdown_to_docx_bytes(md: str, *, title: str | None = None) -> bytes:
    """Render `md` as a DOCX and return the binary content."""
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)
        h.runs[0].font.size = Pt(20)

    for raw_line in md.splitlines():
        line = raw_line.rstrip()

        if not line.strip():
            doc.add_paragraph()
            continue

        if _HR_RE.match(line):
            p = doc.add_paragraph()
            p.add_run("─" * 60).font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)
            continue

        m = _HEADING_RE.match(line)
        if m:
            level = min(len(m.group(1)), 4)
            heading = doc.add_heading(level=level)
            _add_runs(heading, m.group(2))
            continue

        m = _BULLET_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(1))
            continue

        m = _NUMBERED_RE.match(line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(2))
            continue

        p = doc.add_paragraph()
        _add_runs(p, line)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
